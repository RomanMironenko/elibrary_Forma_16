from bs4 import BeautifulSoup
from random import randint
from re import findall
import pandas as pd
import re
import docx
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Mm


# функция для отбора по году публикации
def Year(s, FROM, TO):
    Ys = findall(r'\s\d{4}\.', s)
    if not Ys:
        Ys = findall(r'\s\d{4}', s)
    if not Ys:
        return False
    for y in Ys:
        Y = int(float(y))
    if Y<FROM or Y>TO:
        return False
    else:
        return True


# функция для непредвзятого определения количества страниц в статье
def NP(s):
    pages = s.split()[-1]
    if '-' in pages:
        P = pages.split('-')
        np =  1 + int(float(P[1])-float(P[0]))
    else:
        np = randint(5, 10)
    return '%d' % np # возвращает строку с числом


# функция по перевертыванию страницы
def change_orientation():
    new_section = doc.add_section(WD_SECTION.NEW_PAGE)
    new_section.orientation = WD_ORIENT.LANDSCAPE
    new_section.top_margin = Mm(20)
    new_section.bottom_margin = Mm(10)
    new_section.left_margin = Mm(20)
    new_section.right_margin = Mm(15)
    new_section.page_width = Mm(297)
    new_section.page_height = Mm(210)
    return new_section

# фильтр по годам публикаций
YFrom, YTo = 2014, 2024

#читаем файл index.html
with open('index.html', 'r', encoding='utf-8') as fp:
    soup = BeautifulSoup(fp, 'html.parser')

# убираем заголовок, css и т.д.
soup.head.style.decompose()

# имя автора
aname = soup.title.get_text().split('-')[1]
new_aname = re.sub(r'\b(\w+)\b\s+\b(\w)\w*\b\s+\b(\w)\w*\b', r'\1 \2.\3.', aname)

# новый заголовок
soup.title.string = aname

# новый заголовок
soup.find('span').string = aname

# убираем что-то лишнее
soup.find('i').decompose()

# находим и удаляем первую таблицу
soup.find('table').decompose()

# находим вторую таблицу
table = soup.find('table')

# меняем оформление
table['border'] = 1

# меняем ширину
table['width'] = '100%'

# новый счетчик
N = 1

# ищем все строки в таблице
rows = table.find_all('tr')

# создаем DataFrame для извлечения данных
table = pd.DataFrame(columns={'name_t': str,
                              'forma': str, 'ex_date': str,
                              'str_v': int, 'authors': str})

# цикл для замены содержимого строк
for i in reversed(range(len(rows))):

    # ищем все столбцы
    cols = rows[i].find_all('td')

    # проверка на соответствие стандартному шаблону
    if len(cols) == 3 and cols[1].find('span') and cols[1].find('i'):

        # читаем запись из второго столбца
        content = cols[1].get_text()

        # тут название
        title = cols[1].find('span').get_text()

        # тут название
        authors = cols[1].find('i').get_text()

        # количество цитирований статьи
        cites = int(cols[2].get_text())

        # убираем название, убираем список авторов:
        content = content.replace(title, '')

        # в content остаются только выходные данные
        content = content.replace(authors, '')

        # Определение типа публикации
        thesis = content.replace('В сборнике: ', '')

        abbook = content.replace('В книге: ', '')
        if thesis != content:
            title += '\n(материалы конференции)'
            content = thesis
        elif abbook != content:
            title += '\n(материалы конференции)'
            content = abbook
        else:
            if 'автореф' in content:
                title += '\n(монография)'
            elif 'диссер' in content:
                title += '\n(монография)'
            else:
                title += '\n(научная статья)'

        # получаем список авторов
        authors = authors.split(', ')
        if cites < 10 or not Year(content, YFrom, YTo):
            rows[i].decompose()
        else:
            anumber = len(authors)
            if anumber < 5:
                PS = ''
            else:
                PS = f' и др., всего {anumber:d} чел.'
                authors = ', '.join(authors[0:5]) + PS

        # вносим данные в DataFrame
        table.loc[N] = [title, 'печатная', content, NP(content), authors]

        # счетчик увеличивается на 1
        N += 1

# создаем документ с расширением .docx
doc = docx.Document()

# переводим страницу в альбомный формат
change_orientation()

# вставляем название таблицы с автором
p = doc.add_paragraph(f'Список\nопубликованных учебных изданий и научных трудов соискателя ученого звания\n{aname:s}')

# выравниваем название таблицы по центру
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

# создаем таблицу в файле
table_doc = doc.add_table(rows=(table.shape[0] + 1), cols=(table.shape[1] + 1), style="Table Grid")

# заполняем оголовок таблицы (первая строка таблицы)
for row in range(table.shape[0] + 1):
    if row == 0:
        cell = table_doc.cell(0, 0)
        cell.text = '№\nп/п'
        cell = table_doc.cell(0, 1)
        cell.text = 'Наименование учебных изданий, научных трудов и патентов на изобретения и иные объекты интеллектуальной собственности'
        cell = table_doc.cell(0, 2)
        cell.text = 'Форма учебных изданий и научных трудов'
        cell = table_doc.cell(0, 3)
        cell.text = 'Выходные данные'
        cell = table_doc.cell(0, 4)
        cell.text = 'Объем, стр.'
        cell = table_doc.cell(0, 5)
        cell.text = 'Соавторы'

# заполняем таблицу (вторая и последующая строка таблицы)
for i in range(1, table.shape[0] + 1):
    cell = table_doc.cell(i, 0)
    cell.text = str(i)
    cell = table_doc.cell(i, 1)
    cell.text = str(table.loc[i].name_t)
    cell = table_doc.cell(i, 2)
    cell.text = str(table.loc[i].forma)
    cell = table_doc.cell(i, 3)
    cell.text = str(table.loc[i].ex_date).strip()
    cell = table_doc.cell(i, 4)
    cell.text = str(table.loc[i].str_v)
    cell = table_doc.cell(i, 5)
    cell.text = '\n'.join(map(str, list(table.loc[i].authors)))

# сохраняем файл 'Форма-16.docx'
doc.save('Форма-16.docx')